from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)


def h(level, text):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
    return heading


def p(text="", bold=False, italic=False, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(1)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return para


def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return para


def add_requirement_table(
    code,
    name,
    description,
    actors,
    precondition,
    normal_seq,
    alternate_seq="",
    exceptions="",
    postcondition="",
    comments="",
):
    """Crea una tabla estandarizada para un requisito funcional individual."""
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    rows = [
        ("Código", code, ""),
        ("Nombre", name, ""),
        ("Descripción", description, ""),
        ("Actores", actors, ""),
        ("Precondición", precondition, ""),
        ("Secuencia normal", normal_seq, ""),
    ]
    if alternate_seq:
        rows.append(("Secuencia alterna", alternate_seq, ""))
    if exceptions:
        rows.append(("Excepciones", exceptions, ""))
    if postcondition:
        rows.append(("Postcondición", postcondition, ""))
    if comments:
        rows.append(("Comentarios", comments, ""))
    for label, v1, v2 in rows:
        row = tbl.add_row().cells
        row[0].text = f"**{label}**"
        row[1].text = v1
        row[2].text = v2
    doc.add_paragraph()


# ===================== PORTADA =====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("RUNNERS – PLATAFORMA MÓVIL DE SERVICIOS, DOMICILIOS Y TIENDA LOCAL")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run(
    "Trabajo académico presentado como parte del desarrollo de un proyecto de software"
)
r2.font.size = Pt(12)
r2.font.name = "Calibri"

doc.add_paragraph()
pres = doc.add_paragraph()
pres.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = pres.add_run(
    "Presentado por:\nEquipo de Desarrollo – Runners\n\nUniversidad del Valle\n2025"
)
r3.font.size = Pt(12)
r3.font.name = "Calibri"

doc.add_page_break()

# ===================== 1. RUNNERS =====================
h(1, "1. Runners")
p(
    "Plataforma Móvil Integral para la Gestión de Servicios Locales, Domicilios y Tienda Comunitaria",
    bold=True,
)
p()
p(
    "Runners es una aplicación móvil desarrollada en Flutter que conecta a usuarios de una comunidad con tres servicios clave: una tienda local (comercios y productos), un módulo de domicilios (envíos con asignación automática de repartidores) y un directorio de servicios profesionales (prestadores del hogar, técnicos, profesionales independientes). Además, cuenta con un directorio de contactos de emergencia y utilidad, así como un panel de administración para la gestión global de la plataforma."
)

# ===================== 2. PLANTEAMIENTO =====================
h(1, "2. Planteamiento del problema")
p("Las comunidades locales enfrentan dificultades para:")
for item in [
    "Conectar a sus habitantes con comercios, productos y servicios locales de forma digitalizada.",
    "Gestionar y asignar repartidores de forma eficiente y transparente para servicios de domicilio.",
    "Centralizar en un único canal la oferta de prestadores de servicios profesionales.",
    "Mantener un directorio de contactos de emergencia y utilidad actualizado y accesible.",
    "Ofrecer a los administradores visibilidad sobre operaciones financieras, de servicio y de reparto.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"
p()
p(
    "Sin una solución unificada, la experiencia del usuario se fragmenta: se depende de llamadas telefónicas, grupos de mensajería informal y procesos manuales que generan errores, retrasos e informalidad en las transacciones locales."
)

# ===================== 3. JUSTIFICACIÓN =====================
h(1, "3. Justificación")
p(
    "Este proyecto busca digitalizar la economía local a través de una plataforma móvil integral. Con ello, se busca:"
)
for item in [
    "Mejorar la experiencia del cliente, brindando acceso rápido a tiendas, servicios y domicilios.",
    "Incrementar la competitividad de los prestadores y comercios locales.",
    "Optimizar los procesos de asignación de domicilios mediante un sistema automático.",
    "Formalizar la oferta de servicios profesionales con validación administrativa.",
    "Centralizar la información operativa en un panel administrativo con reportes en tiempo real.",
    "Reducir errores humanos al automatizar flujos de solicitud, aprobación y seguimiento.",
    "Fomentar la confianza entre usuarios mediante perfiles verificados y seguimiento transparente.",
    "Escalar y extender la plataforma fácilmente gracias a Clean Architecture.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 4. OBJETIVO GENERAL =====================
h(1, "4. Objetivo general")
p(
    "Desarrollar una aplicación móvil integral para comunidades locales que proporcione una experiencia interactiva, moderna y eficiente, permitiendo a los usuarios acceder a servicios de tienda, domicilios y prestadores profesionales desde una única plataforma, con autenticación segura basada en roles, gestión de estado en tiempo real y panel administrativo completo."
)
p()
p(
    "El sistema incorpora autenticación JWT con roles diferenciados (CLIENTE, PRESTADOR, DOMICILIARIO, ADMIN), módulo de tienda con carrito Hive, módulo de domicilios con asignación automática de repartidores, módulo de servicios con validación de prestadores, directorio de contactos y dashboard de administración con reportes."
)
p()
p(
    "Arquitectura técnica: backend Django REST Framework + JWT; frontend Flutter con Clean Architecture, Riverpod, GoRouter y Dio."
)

h(2, "4.1. Objetivos específicos")
for item in [
    "Diseñar y desarrollar un backend robusto con Django y DRF para los módulos de usuarios, tienda, domicilios, servicios, contactos y reportes.",
    "Implementar un frontend móvil en Flutter que consuma los endpoints del backend mediante Dio con interceptores JWT.",
    "Desarrollar autenticación con JWT, soporte para roles múltiples y refresco automático de tokens.",
    "Diseñar un modelo de base de datos relacional optimizado para el dominio de servicios locales.",
    "Implementar asignación automática de domiciliarios disponibles a solicitudes entrantes.",
    "Desarrollar flujo completo de compra en tienda con carrito Hive y seguimiento de órdenes.",
    "Implementar validación y aprobación de prestadores de servicios por el administrador.",
    "Desarrollar panel de administración con dashboard, reportes y gestión global.",
    "Aplicar metodología ágil con historias de usuario, sprints y control de versiones en GitHub.",
    "Realizar pruebas integrales en backend (Postman) y frontend (flujos completos).",
    "Documentar exhaustivamente arquitectura, modelo relacional, diagramas y guías de instalación.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 5. ANTECEDENTES =====================
h(1, "5. Antecedentes")
p("Plataformas de servicios locales y domicilios (Rappi, Domicilios.com, TaskRabbit) muestran que:")
for item in [
    "La digitalización de servicios locales reduce intermediación y mejora la transparencia.",
    "Los sistemas de asignación automática de repartidores reducen tiempos de espera.",
    "Los directorios de prestadores verificados generan mayor confianza y uso recurrente.",
    "Flutter/Dart y Django/DRF son stacks reconocidos por productividad y escalabilidad.",
    "JWT es el estándar para autenticación sin estado en APIs REST.",
    "Clean Architecture en Flutter permite mantenimiento ágil y pruebas efectivas.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 6. MARCO LEGAL =====================
h(1, "6. Marco legal")
for item in [
    "Protección de datos personales: Ley 1581 de 2012 (Habeas Data, Colombia). Minimización y consentimiento.",
    "Políticas y Términos: textos legales de uso del servicio y privacidad antes del despliegue público.",
    "Derechos de autor: imágenes propias o con licencia libre. Documentos de prestadores son su responsabilidad.",
    "Seguridad: contraseñas con hash bcrypt (Django), tokens en FlutterSecureStorage, sin datos bancarios.",
    "Nota: el módulo de pagos es simulado en la fase actual. Integración real requiere evaluación legal adicional.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 7. MARCO TEÓRICO =====================
h(1, "7. Marco teórico")
for item in [
    "Arquitectura cliente/servidor: el frontend Flutter consume el backend RESTful Django/DRF.",
    "REST & JSON: diseño de endpoints con recursos y verbos HTTP estándar (GET, POST, PUT, PATCH, DELETE).",
    "JWT: flujo access/refresh para autenticación stateless; interceptor Dio renueva tokens automáticamente ante 401.",
    "Clean Architecture (Flutter): capas domain (entidades, casos de uso), data (repositorios, datasources, modelos) y presentation (providers, pantallas, widgets).",
    "Riverpod: gestión de estado reactiva con StateNotifierProvider y FutureProvider.",
    "GoRouter: enrutamiento declarativo con guardias de rol.",
    "Hive: persistencia local offline del carrito de compras.",
    "Modelo relacional: tablas normalizadas, relaciones 1:N y N:M con tablas intermedias en Django ORM.",
    "MVC en Django: modelos (ORM), ViewSets (DRF), serializadores, URLs.",
    "UX centrado en el usuario: loading states, empty states, error messages, navegación por roles.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 8. REQUERIMIENTOS =====================
h(1, "8. Requerimientos funcionales y no funcionales")
h(2, "8.1 Introducción")
p(
    "En esta sección se describe el desarrollo del sistema Runners, una plataforma móvil diseñada para conectar a usuarios de una comunidad con servicios de tienda, domicilios y profesionales independientes. El sistema gestiona cuatro tipos de usuarios (CLIENTE, PRESTADOR, DOMICILIARIO, ADMIN) con flujos diferenciados por rol."
)

h(2, "8.2 Requisitos Funcionales a Nivel del Sistema")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "ID"
hdr[1].text = "Módulo"
hdr[2].text = "Requisito"
reqs = [
    ("REQ-01", "Autenticación", "Registro de usuario con nombre, correo, contraseña y rol"),
    ("REQ-02", "Autenticación", "Inicio de sesión con JWT (access + refresh token)"),
    ("REQ-03", "Tienda", "Listado de comercios y productos por categoría"),
    ("REQ-04", "Tienda", "Carrito de compras persistido localmente (Hive)"),
    ("REQ-05", "Tienda", "Creación y seguimiento de órdenes de compra"),
    ("REQ-06", "Domicilios", "Solicitud de domicilio con dirección origen y destino"),
    ("REQ-07", "Domicilios", "Asignación automática de domiciliario disponible"),
    ("REQ-08", "Domicilios", "Seguimiento del estado del domicilio"),
    ("REQ-09", "Servicios", "Registro de prestadores con perfil, foto y hoja de vida"),
    ("REQ-10", "Servicios", "Aprobación/rechazo de prestadores por el administrador"),
    ("REQ-11", "Servicios", "Solicitud de servicio por parte del cliente"),
    ("REQ-12", "Contactos", "Directorio de contactos de emergencia y utilidad"),
    ("REQ-13", "Admin", "Dashboard con métricas de operación"),
    ("REQ-14", "Admin", "Reportes de domicilios, servicios y tienda"),
    ("REQ-15", "Admin", "Gestión de usuarios, comercios y configuración"),
    ("REQ-16", "Domicilios", "Registro y gestión de domiciliarios"),
    ("REQ-17", "Domicilios", "Registro de ingresos y egresos de domiciliarios"),
    ("REQ-18", "Autenticación", "Cierre de sesión seguro para invalidar tokens activos"),
    ("REQ-19", "Autenticación", "Roles diferenciados y control de acceso por perfil"),
    ("REQ-20", "Servicios", "Búsqueda de prestadores por categoría de servicio"),
    ("REQ-21", "Servicios", "Cambio de estado del prestador (DISPONIBLE/OCUPADO/INACTIVO)"),
    ("REQ-22", "Domicilios", "Cambio de estado del domiciliario (DISPONIBLE/OCUPADO/INACTIVO)"),
    ("REQ-23", "Domicilios", "Solicitud directa de domicilio a un domiciliario disponible"),
    ("REQ-24", "Contactos", "Filtrado de contactos por estado de disponibilidad"),
    ("REQ-25", "Admin", "Reporte financiero agregado de domiciliarios"),
    ("REQ-26", "Infraestructura", "Uso de variables de entorno (.env) para configuraciones sensibles"),
    ("REQ-27", "Infraestructura", "PostgreSQL como base de datos de producción"),
    ("REQ-28", "Seguridad API", "Configuración de CORS para el backend de Runners"),
]
for r in reqs:
    row = tbl.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
doc.add_paragraph()

h(2, "8.3 REQ-01 — Registro de usuario")
add_requirement_table(
    code="REQ-01",
    name="Registro de usuario",
    description=(
        "Permite que un nuevo usuario cree una cuenta proporcionando nombre, correo y contraseña, "
        "quedando registrado por defecto como CLIENTE (o como el rol permitido que seleccione)."
    ),
    actors="Usuario (cliente, prestador, domiciliario, admin), App (Flutter/Web), API Django",
    precondition="El usuario no debe estar registrado previamente con el mismo correo electrónico.",
    normal_seq=(
        "El usuario diligencia el formulario de registro, la app envía la solicitud al backend, "
        "el backend valida datos y unicidad del correo y, si todo es correcto, crea el usuario y responde éxito."
    ),
    alternate_seq=(
        "Si el correo ya existe, el backend rechaza el registro y la app muestra un mensaje indicando que el correo ya está registrado."
    ),
    exceptions=(
        "Errores de conexión o datos malformados se responden con códigos 4xx/5xx y mensajes estándar de error."
    ),
    postcondition=(
        "El usuario queda creado en la base de datos con rol y estado inicial definidos y puede continuar al flujo de inicio de sesión (REQ-02)."
    ),
    comments=(
        "La contraseña se almacena con hash seguro de Django y el correo se define como único en el modelo de usuario."
    ),
)

h(2, "8.4 REQ-02 — Inicio de sesión y gestión de tokens JWT")
add_requirement_table(
    code="REQ-02",
    name="Inicio de sesión y gestión de tokens JWT",
    description=(
        "Permite a los usuarios autenticarse con correo y contraseña para obtener tokens de acceso y refresco, "
        "gestionando la renovación automática mientras la sesión siga vigente."
    ),
    actors="Cliente, Prestador, Domiciliario, Admin, App Flutter/Web, API Django",
    precondition="El usuario debe estar registrado y activo en el sistema.",
    normal_seq=(
        "El usuario ingresa sus credenciales, la app envía la petición de login, el backend valida y retorna access_token y refresh_token, "
        "la app almacena los tokens de forma segura, consulta el perfil y redirige según el rol."
    ),
    alternate_seq=(
        "Si las credenciales son incorrectas o el usuario está suspendido, el backend rechaza la solicitud y la app muestra el mensaje apropiado."
    ),
    exceptions=(
        "Si el access_token expira, el interceptor intenta renovarlo con el refresh_token; si este también falla, se cierra la sesión y se redirige al login."
    ),
    postcondition=(
        "El usuario queda autenticado con tokens válidos almacenados y su perfil cargado en la app; el backend reconoce al usuario en peticiones subsecuentes."
    ),
    comments=(
        "Se usa djangorestframework-simplejwt en backend y se recomiendan interceptores y guards de ruta en frontend."
    ),
)

h(2, "8.5 REQ-03 — Visualización de tienda por categorías")
add_requirement_table(
    code="REQ-03",
    name="Visualización de tienda por categorías",
    description=(
        "Permite al cliente explorar la tienda, filtrar por categorías y ver el catálogo de comercios y productos disponibles con su información básica."
    ),
    actors="Cliente, App Flutter/Web, API Django",
    precondition=(
        "El cliente debe estar autenticado y deben existir categorías, comercios y productos registrados y activos."
    ),
    normal_seq=(
        "El cliente ingresa al módulo de tienda, la app consulta categorías y comercios, muestra los productos por categoría y permite seleccionar productos para el carrito."
    ),
)

h(2, "8.6 REQ-04/REQ-05 — Carrito y orden de compra")
add_requirement_table(
    code="REQ-04 / REQ-05",
    name="Carrito de compras y generación de orden",
    description=(
        "Permite al cliente agregar productos al carrito, ajustar cantidades y confirmar un pedido, "
        "generando una orden en el backend con estado inicial PENDIENTE."
    ),
    actors="Cliente, App Flutter/Web, API Django, Comercio",
    precondition=(
        "El cliente debe estar autenticado, haber seleccionado productos válidos y existir productos disponibles asociados a un comercio."
    ),
    normal_seq=(
        "El cliente agrega productos al carrito, revisa el resumen, confirma la compra y la app envía la orden al backend, que valida y crea la Orden y sus ítems."
    ),
    alternate_seq="El cliente puede vaciar o modificar el carrito antes de confirmar la orden.",
    exceptions=(
        "Si hay productos inactivos o datos inconsistentes en el carrito, el backend rechaza la orden con mensajes de validación."
    ),
)

h(2, "8.7 REQ-06/REQ-07 — Solicitud de domicilio y asignación automática")
add_requirement_table(
    code="REQ-06 / REQ-07",
    name="Solicitud de domicilio y asignación automática",
    description=(
        "Permite al cliente solicitar un domicilio indicando datos de recogida y entrega, mientras el sistema asigna automáticamente un domiciliario disponible siguiendo una regla de negocio definida."
    ),
    actors="Cliente, Domiciliario, App Flutter/Web, API Django",
    precondition=(
        "El cliente debe estar autenticado y debe existir al menos un domiciliario con estado DISPONIBLE."
    ),
    normal_seq=(
        "El cliente crea la solicitud de domicilio, el backend la registra y ejecuta la lógica de asignación automática para seleccionar un domiciliario disponible y actualizar el estado."
    ),
    alternate_seq=(
        "Si no hay domiciliarios DISPONIBLE, la solicitud puede quedar en espera y se notifica al cliente de la falta de disponibilidad."
    ),
)

h(2, "8.8 REQ-08 — Seguimiento de domicilio")
add_requirement_table(
    code="REQ-08",
    name="Seguimiento del estado del domicilio",
    description=(
        "Permite a cliente y domiciliario consultar y actualizar el estado de una solicitud de domicilio a lo largo de su ciclo de vida."
    ),
    actors="Cliente, Domiciliario, Administrador, App Flutter/Web, API Django",
    precondition="Debe existir una solicitud de domicilio registrada y los usuarios deben estar autenticados.",
    normal_seq=(
        "El cliente consulta el detalle y ve el estado, el domiciliario actualiza el estado según avanza la entrega y el backend registra cada cambio mediante endpoints dedicados."
    ),
)

h(2, "8.9 REQ-09 — Perfil de prestador de servicios")
add_requirement_table(
    code="REQ-09",
    name="Gestión de perfil de prestador de servicios",
    description=(
        "Permite a un usuario con rol PRESTADOR completar y actualizar su perfil profesional, incluyendo foto, descripción, categoría y hoja de vida."
    ),
    actors="Prestador de Servicio, Administrador, App Flutter/Web, API Django",
    precondition="El usuario debe estar registrado con rol PRESTADOR y autenticado en la app.",
    normal_seq=(
        "El prestador accede a su perfil, completa o actualiza la información y adjunta archivos; la app envía los datos y el backend guarda el perfil en estado PENDIENTE."
    ),
)

h(2, "8.10 REQ-10 — Aprobación de prestadores")
add_requirement_table(
    code="REQ-10",
    name="Aprobación o rechazo de prestadores de servicios",
    description=(
        "Permite al administrador revisar los perfiles de prestadores y decidir si aprueba o rechaza su participación en la plataforma."
    ),
    actors="Administrador, Prestador de Servicio, API Django, App Flutter/Web",
    precondition=(
        "Deben existir perfiles de prestadores en estado PENDIENTE y el administrador debe estar autenticado con rol ADMIN."
    ),
    normal_seq=(
        "El administrador revisa los perfiles pendientes, decide aprobar o rechazar y el backend actualiza el estado y registra el motivo si aplica."
    ),
)

h(2, "8.11 REQ-11 — Solicitud de servicio")
add_requirement_table(
    code="REQ-11",
    name="Solicitud de servicio a prestador aprobado",
    description=(
        "Permite al cliente solicitar un servicio a un prestador aprobado, definiendo los detalles básicos y permitiendo su aceptación o rechazo por parte del prestador."
    ),
    actors="Cliente, Prestador de Servicio, App Flutter/Web, API Django",
    precondition=(
        "Debe existir al menos un prestador en estado APROBADO y tanto cliente como prestador deben estar autenticados."
    ),
    normal_seq=(
        "El cliente crea la solicitud hacia un prestador aprobado, el backend la guarda en estado PENDIENTE y el prestador puede aceptarla o rechazarla desde su panel."
    ),
)

h(2, "8.12 REQ-12 — Directorio de contactos")
add_requirement_table(
    code="REQ-12",
    name="Directorio de contactos de emergencia y utilidad",
    description=(
        "Permite a los usuarios consultar un directorio de contactos relevantes (emergencias, profesionales, comercios) y al administrador gestionar su contenido."
    ),
    actors="Cliente, Prestador, Domiciliario, Admin, API Django, App Flutter/Web",
    precondition="Deben existir contactos registrados y los usuarios deben estar autenticados según políticas definidas.",
    normal_seq=(
        "El usuario accede al módulo de contactos, la app consulta y muestra la lista y permite aplicar filtros o búsquedas; el administrador puede crear, editar o desactivar contactos."
    ),
)

h(2, "8.13 REQ-13 — Dashboard administrativo")
add_requirement_table(
    code="REQ-13",
    name="Panel de métricas administrativas",
    description=(
        "Proporciona al administrador un panel con indicadores clave sobre usuarios, órdenes, domicilios y servicios para el monitoreo operativo."
    ),
    actors="Administrador, API Django, App Flutter/Web",
    precondition="El administrador debe estar autenticado y deben existir datos registrados en el sistema.",
    normal_seq=(
        "El administrador abre el dashboard, la app consulta el endpoint de resumen y el backend responde con métricas agregadas que se muestran en tarjetas y gráficos."
    ),
)

h(2, "8.14 REQ-14 — Reportes de operación")
add_requirement_table(
    code="REQ-14",
    name="Generación de reportes de tienda, servicios y domicilios",
    description=(
        "Permite al administrador consultar reportes filtrados sobre domicilios realizados, servicios prestados y ventas de tienda para periodos definidos."
    ),
    actors="Administrador, API Django, App Flutter/Web",
    precondition="El administrador debe estar autenticado y debe existir información histórica suficiente.",
    normal_seq=(
        "El administrador selecciona tipo de reporte y filtros, la app envía la consulta y el backend ejecuta agregaciones y devuelve los datos para mostrarlos en tablas o gráficos."
    ),
)

h(2, "8.15 REQ-15 — Gestión administrativa")
add_requirement_table(
    code="REQ-15",
    name="Gestión administrativa de usuarios, comercios y configuración",
    description=(
        "Permite al administrador gestionar usuarios, comercios, parámetros de configuración y registros financieros asociados a domiciliarios."
    ),
    actors="Administrador, API Django, App Flutter/Web",
    precondition="El administrador debe estar autenticado con rol ADMIN.",
    normal_seq=(
        "El administrador accede al módulo de administración, gestiona altas y bajas lógicas de usuarios y comercios y ajusta parámetros clave del sistema."
    ),
)

h(2, "8.16 REQ-16 — Registro y gestión de domiciliarios")
add_requirement_table(
    code="REQ-16",
    name="Registro y gestión de domiciliarios",
    description=(
        "Permite registrar domiciliarios con sus datos básicos y gestionar su estado (DISPONIBLE/OCUPADO/INACTIVO), de modo que puedan participar en el módulo de domicilios."
    ),
    actors="Administrador, Domiciliario, API Django, App (panel de domiciliarios/web)",
    precondition=(
        "El administrador debe estar autenticado con rol ADMIN y debe existir o crearse un usuario base asociado al domiciliario."
    ),
    normal_seq=(
        "El administrador registra al domiciliario con sus datos, el backend valida unicidad del número asignado y crea el registro con estado DISPONIBLE e is_active=True."
    ),
)

h(2, "8.17 REQ-17 — Registro de ingresos y egresos de domiciliarios")
add_requirement_table(
    code="REQ-17",
    name="Registro de ingresos y egresos de domiciliarios",
    description=(
        "Permite a los domiciliarios registrar ingresos por servicios realizados y egresos (gastos), para llevar control de su balance y de la comisión que corresponde a Runners."
    ),
    actors="Domiciliario, Administrador, API Django, App (panel de domiciliarios/web)",
    precondition="Debe existir un domiciliario registrado asociado al usuario autenticado.",
    normal_seq=(
        "El domiciliario accede a su panel financiero, registra ingresos o egresos y el backend guarda FinancialRecord y actualiza los totales calculados."
    ),
)

h(2, "8.18 REQ-18 — Cierre de sesión seguro")
add_requirement_table(
    code="REQ-18",
    name="Cierre de sesión seguro",
    description=(
        "Permite al usuario cerrar sesión de forma segura, eliminando los tokens almacenados y evitando que permanezcan activos en el dispositivo."
    ),
    actors="Usuario autenticado, App Flutter/Web",
    precondition="El usuario debe estar autenticado y tener tokens activos almacenados en el cliente.",
    normal_seq=(
        "El usuario selecciona la opción de cerrar sesión, la app limpia los tokens del almacenamiento seguro y redirige a la pantalla de login."
    ),
)

h(2, "8.19 REQ-19 — Roles diferenciados y control de acceso")
add_requirement_table(
    code="REQ-19",
    name="Roles diferenciados y control de acceso",
    description=(
        "Define y aplica roles de usuario (CLIENTE, PRESTADOR, DOMICILIARIO, ADMIN) para controlar el acceso a pantallas y endpoints."
    ),
    actors="Todos los usuarios, Backend Django, Frontend Flutter/Web",
    precondition="El usuario debe estar autenticado con un rol válido.",
    normal_seq=(
        "El backend incluye el rol en el token o perfil, el frontend lee el rol y muestra sólo las pantallas permitidas, mientras el backend aplica permisos por rol en cada endpoint."
    ),
)

h(2, "8.20 REQ-20 — Búsqueda de prestadores por categoría")
add_requirement_table(
    code="REQ-20",
    name="Búsqueda de prestadores por categoría",
    description=(
        "Permite a los clientes buscar prestadores de servicios filtrando por categoría para encontrar al profesional adecuado."
    ),
    actors="Cliente, App Flutter/Web, API Django",
    precondition="Debe existir al menos un prestador aprobado asociado a categorías de servicios.",
    normal_seq=(
        "El cliente selecciona una categoría o aplica filtros, la app consulta el backend y muestra la lista de prestadores aprobados que cumplen los criterios."
    ),
)

h(2, "8.21 REQ-21 — Cambio de estado del prestador")
add_requirement_table(
    code="REQ-21",
    name="Cambio de estado del prestador",
    description=(
        "Permite actualizar el estado del prestador (DISPONIBLE/OCUPADO/INACTIVO) para reflejar su disponibilidad en el módulo de servicios."
    ),
    actors="Prestador, Administrador, API Django, App Flutter/Web",
    precondition="El prestador debe existir y estar aprobado en la plataforma.",
    normal_seq=(
        "El prestador o el administrador cambia el estado desde su panel y el backend actualiza el registro para afectar su visibilidad y disponibilidad en las búsquedas."
    ),
)

h(2, "8.22 REQ-22 — Cambio de estado del domiciliario")
add_requirement_table(
    code="REQ-22",
    name="Cambio de estado del domiciliario",
    description=(
        "Permite actualizar el estado del domiciliario (DISPONIBLE/OCUPADO/INACTIVO) para controlar la asignación automática de domicilios."
    ),
    actors="Domiciliario, Administrador, API Django, App Flutter/Web",
    precondition="El domiciliario debe estar registrado y activo.",
    normal_seq=(
        "El domiciliario o el administrador modifica el estado y el backend actualiza el registro, afectando su inclusión en la lógica de asignación automática."
    ),
)

h(2, "8.23 REQ-23 — Solicitud directa de domicilio a un domiciliario disponible")
add_requirement_table(
    code="REQ-23",
    name="Solicitud directa de domicilio a domiciliario disponible",
    description=(
        "Permite que el administrador o un flujo específico envíe una solicitud de domicilio dirigida a un domiciliario disponible en particular."
    ),
    actors="Administrador, Domiciliario, API Django, App Flutter/Web",
    precondition="Debe existir un domiciliario en estado DISPONIBLE y un cliente o solicitud que requiera servicio.",
    normal_seq=(
        "Se selecciona un domiciliario disponible, se crea la solicitud de domicilio asociada a él y el backend registra el vínculo y el estado inicial correspondiente."
    ),
)

h(2, "8.24 REQ-24 — Filtrado de contactos por disponibilidad")
add_requirement_table(
    code="REQ-24",
    name="Filtrado de contactos por disponibilidad",
    description=(
        "Permite filtrar el directorio de contactos por estado de disponibilidad o tipo para facilitar la búsqueda de información útil."
    ),
    actors="Usuario, App Flutter/Web, API Django",
    precondition="Deben existir contactos con distintos tipos o estados registrados.",
    normal_seq=(
        "El usuario aplica filtros sobre el directorio y la app consulta el backend o filtra localmente para mostrar sólo los contactos que cumplen los criterios."
    ),
)

h(2, "8.25 REQ-25 — Reporte financiero agregado de domiciliarios")
add_requirement_table(
    code="REQ-25",
    name="Reporte financiero agregado de domiciliarios",
    description=(
        "Permite al administrador consultar un reporte agregado de ingresos, egresos y comisiones de los domiciliarios."
    ),
    actors="Administrador, API Django, App Flutter/Web",
    precondition="Deben existir registros financieros asociados a domiciliarios.",
    normal_seq=(
        "El administrador accede al módulo de reportes financieros, selecciona filtros de fecha y el backend agrega los FinancialRecord por domiciliario y retorna los totales."
    ),
)

h(2, "8.26 REQ-26 — Uso de variables de entorno (.env)")
add_requirement_table(
    code="REQ-26",
    name="Uso de variables de entorno (.env)",
    description=(
        "Establece el uso de variables de entorno para credenciales, claves secretas y configuraciones sensibles tanto en backend como en frontend."
    ),
    actors="Equipo de despliegue, Backend Django, Frontend Flutter/Web",
    precondition="Debe existir un archivo .env o mecanismo equivalente configurado en cada entorno.",
    normal_seq=(
        "Durante la configuración, se definen las variables necesarias (SECRET_KEY, credenciales de BD, URLs, etc.) y la aplicación las lee sin exponer valores sensibles en el código fuente."
    ),
)

h(2, "8.27 REQ-27 — Configuración de PostgreSQL como base de datos de producción")
add_requirement_table(
    code="REQ-27",
    name="Configuración de PostgreSQL como base de datos de producción",
    description=(
        "Define que, para entornos de producción, la base de datos relacional utilizada será PostgreSQL, reemplazando SQLite del entorno de desarrollo."
    ),
    actors="Equipo de infraestructura, Backend Django",
    precondition="Debe existir una instancia de PostgreSQL accesible y credenciales configuradas en variables de entorno.",
    normal_seq=(
        "En el entorno de producción se configura Django para usar PostgreSQL, se ejecutan migraciones y se verifica el correcto funcionamiento de la API contra la nueva base de datos."
    ),
)

h(2, "8.28 REQ-28 — Configuración de CORS para el backend de Runners")
add_requirement_table(
    code="REQ-28",
    name="Configuración de CORS para el backend de Runners",
    description=(
        "Establece la configuración de CORS en el backend para permitir que los frontends autorizados (web/móvil) consuman la API de forma segura."
    ),
    actors="Equipo de infraestructura, Backend Django",
    precondition="El proyecto debe tener instalado y configurado django-cors-headers u otra solución equivalente.",
    normal_seq=(
        "Se definen los orígenes permitidos y encabezados necesarios para las peticiones desde los clientes, restringiendo el acceso desde dominios no autorizados."
    ),
)

h(2, "8.29 Requerimientos no funcionales")
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0].cells
hdr2[0].text = "ID"
hdr2[1].text = "Categoría"
hdr2[2].text = "Descripción"
rnfs = [
    (
        "RNF-01",
        "Seguridad",
        "Contraseñas hash bcrypt, JWT para sesiones, FlutterSecureStorage para tokens y uso de variables de entorno (.env) para credenciales y secretos.",
    ),
    ("RNF-02", "Rendimiento", "Latencia < 500ms; paginación en listados extensos"),
    ("RNF-03", "Escalabilidad", "Clean Architecture; compatible con Docker"),
    (
        "RNF-04",
        "Mantenibilidad",
        "Apps Django por dominio; features Flutter por módulo",
    ),
    (
        "RNF-05",
        "Observabilidad",
        "Logs backend; manejo de errores estructurado en frontend",
    ),
    (
        "RNF-06",
        "Compatibilidad",
        "Aplicación móvil Android/iOS; backend compatible con PostgreSQL en producción",
    ),
    (
        "RNF-07",
        "Disponibilidad",
        "Backend 24/7 en producción con gunicorn/supervisor",
    ),
    (
        "RNF-08",
        "Infraestructura",
        "Uso de PostgreSQL como base de datos de producción para garantizar escalabilidad, rendimiento y confiabilidad.",
    ),
    (
        "RNF-09",
        "Seguridad API",
        "Configuración adecuada de CORS en el backend para permitir que el frontend (web/móvil) se comunique con la API sin exponer el sistema a orígenes no autorizados.",
    ),
]
for r in rnfs:
    row = tbl2.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
doc.add_paragraph()

# ===================== 9. CUALIDADES =====================
h(1, "9. Cualidades del Sistema")
h(2, "9.1 Usabilidad y Accesibilidad")
p(
    "Navegación por roles (GoRouter), estados visuales (AppLoading, AppError, AppEmptyState), widgets compartidos (AppButton, AppTextField, AppCard), formularios con validación en tiempo real y mensajes de error descriptivos."
)
h(3, "9.2 Confiabilidad")
p(
    "Reintentos automáticos con interceptor Dio, persistencia del carrito en Hive ante desconexiones, renovación automática de tokens JWT."
)
h(4, "9.3 Rendimiento")
p(
    "ORM con select_related/prefetch_related; paginación en listados; CachedNetworkImage en Flutter; Hive sin latencia de red para carrito."
)
h(5, "9.4 Soportabilidad")
p(
    "Clean Architecture modular: cada feature Flutter y cada app Django es independiente. Nuevos módulos se agregan sin modificar los existentes."
)

# ===================== 10. INTERFACES =====================
h(1, "10. Interfaces del Sistema")
h(2, "10.1 Interfaces de Usuario")
h(3, "10.1.1 Aspecto y Sensación")
p(
    "Diseño moderno y coherente con tema definido en core/theme/. Iconos Material Design. Paleta de colores unificada."
)
h(4, "10.1.2 Diseño y Navegación")
p(
    "BottomNavigationBar o NavigationDrawer con accesos según rol. Pantalla Splash gestiona redirección inicial basada en token almacenado."
)
h(5, "10.1.3 Consistencia")
p(
    "ThemeData centralizado en core/theme/app_theme.dart. AppButton y AppTextField garantizan homogeneidad."
)
h(6, "10.1.4 Personalización por Rol")
for rol, accesos in [
    ("CLIENTE", "tienda, carrito, órdenes, domicilios, servicios (solicitar), contactos, perfil"),
    ("PRESTADOR", "perfil, mis solicitudes de servicio, disponibilidad"),
    (
        "DOMICILIARIO",
        "solicitudes asignadas, cambio de estado, historial financiero",
    ),
    ("ADMIN", "dashboard, usuarios, prestadores, reportes, contactos, configuración"),
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(f"{rol}: {accesos}").font.name = "Calibri"

h(2, "10.2 Interfaces con Sistemas Externos")
h(3, "10.2.1 Interfaces de Software")
p(
    "Flutter/Dart (SDK 3.9+) — API REST Django/DRF (Python 3.12, Django 5.2.6). Comunicación JSON sobre HTTP/HTTPS."
)
p("10.2.2 Interfaces de Hardware")
p(
    "Sin hardware especializado. Android 5.0+ / iOS 12+. Servidor: Linux/Windows con Python 3.12+."
)
h(5, "10.2.3 Interfaces de Comunicaciones")
p(
    "HTTPS en producción. Authorization: Bearer <token> en cada petición protegida. HTTP sobre localhost:8000 en desarrollo."
)

# ===================== 11. REGLAS DE NEGOCIO =====================
h(1, "11. Reglas de Negocio")
h(2, "11.1 Usuarios y Autenticación")
h(3, "11.1.1 RN-01 – Roles Exclusivos")
p(
    "Si un usuario se registra con rol CLIENTE, no podrá acceder a funciones de DOMICILIARIO, PRESTADOR ni ADMIN. Guardias de rol en GoRouter y permisos en backend."
)
h(4, "11.1.2 RN-02 – Validación de Registro")
p(
    "Si el correo ya existe, el sistema notifica el conflicto e impide duplicados. email es unique=True en modelo User."
)
h(5, "11.1.3 RN-03 – Inicio de Sesión Seguro")
p(
    "Validación con hash bcrypt. Token de acceso expira en 5 minutos; refresh en 1 día."
)
h(2, "11.2 Tienda y Carrito")
h(3, "11.2.1 RN-04 – Validación del Carrito")
p(
    "Si no hay ítems con cantidad > 0, el sistema no permite confirmar la compra."
)
h(4, "11.2.2 RN-05 – Disponibilidad de Productos")
p(
    "Si is_available = False, el producto no aparece en listado ni puede agregarse al carrito."
)
h(5, "11.2.3 RN-06 – Generación de Orden")
p(
    "Al confirmar el carrito, se crea una Orden con estado PENDIENTE y OrderItems con precio histórico."
)
h(2, "11.3 Domicilios")
h(3, "11.3.1 RN-07 – Asignación Automática")
p(
    "Al crear solicitud, el sistema asigna el domiciliario DISPONIBLE de menor assigned_number. Si no hay disponibles, deliverer = null."
)
h(4, "11.3.2 RN-08 – Cambio de Estado a ENTREGADO")
p(
    "Se registra completed_at, se libera al domiciliario (DISPONIBLE) y se genera FinancialRecord."
)
h(5, "11.3.3 RN-09 – Registro Financiero")
p(
    "runners_commission = delivery_fee x commission_rate (de SystemConfig)."
)
h(2, "11.4 Servicios")
h(3, "11.4.1 RN-10 – Validación de Prestador")
p(
    "Si approval_status != APROBADO, el prestador no aparece en el directorio de servicios."
)
h(4, "11.4.2 RN-11 – Solicitud Única")
p(
    "Si hay solicitud PENDIENTE o APROBADA con el mismo prestador, el sistema impide duplicados."
)
h(5, "11.4.3 RN-12 – Cálculo de Tarifas")
p(
    "client_total = provider_fee + runners_fee (runners_fee según SystemConfig)."
)
h(2, "11.5 Contactos")
h(3, "11.5.1 RN-13 – Visibilidad")
p(
    "Solo contactos con is_active = True son visibles en el directorio público."
)
h(4, "11.5.2 RN-14 – Gestión Exclusiva")
p(
    "Solo ADMIN puede crear, editar o desactivar contactos."
)

# ===================== 12. RESTRICCIONES =====================
h(1, "12. Restricciones del Sistema")
h(2, "12.1 Lenguaje de Implementación")
p(
    "Backend: Django (Python 3.12) con DRF. Frontend: Flutter (Dart 3.9+) con Clean Architecture y Riverpod."
)
h(3, "12.2 Base de Datos")
p(
    "Desarrollo: SQLite. Producción: PostgreSQL con psycopg2-binary 2.9.10."
)
h(4, "12.3 Herramientas de Desarrollo")
herramientas = [
    "Python 3.12, Dart 3.9+ / Flutter 3.9+",
    "Django 5.2.6, Django REST Framework 3.16.1",
    "djangorestframework-simplejwt 5.5.1, django-cors-headers 4.8.0, Pillow 11.3.0",
    "flutter_riverpod 2.x, go_router 13.x, dio 5.x, hive 2.x, flutter_secure_storage 9.x",
    "Git + GitHub (https://github.com/Karatsuyu/Runners.git)",
    "Postman – coleccion runners_api_postman.json (30+ endpoints)",
    "VS Code con extensiones Flutter y Python",
]
for item in herramientas:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"
h(4, "12.4 Límites de Recursos")
p(
    "Dispositivo movil: minimo 2 GB RAM. Servidor desarrollo: minimo 1 GB RAM, 1 nucleo. Produccion recomendada: 2 GB RAM, 2 nucleos."
)

# ===================== 13. DOCUMENTACIÓN =====================
h(1, "13. Documentación del Sistema")
h(2, "13.1 Sistema de Ayuda")
p(
    "Mensajes contextuales, indicadores de estado (cargando, error, vacío) y errores del API en lenguaje natural."
)
h(3, "13.2 Documentación Técnica")
for item in [
    "Diagrama ER del modelo relacional (Mermaid).",
    "Diagrama de arquitectura del sistema.",
    "Coleccion Postman con 30+ endpoints (runners_api_postman.json).",
    "README.md, QUICKSTART.md con instrucciones de instalacion.",
    "Guias de implementacion Flutter y web.",
    "CONTRIBUTING.md y PR template en español.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"
h(4, "13.3 Responsabilidad")
p(
    "El equipo de desarrollo es responsable de la creacion y mantenimiento de la documentacion, actualizandola con cada cambio significativo."
)

# ===================== 14. MARCO METODOLÓGICO =====================
h(1, "14. Marco metodológico")
h(2, "14.1 Metodología de proyecto")
p(
    "Metodologia agil (Scrum/Kanban): sprints de 1-2 semanas, historias de usuario organizadas por modulo, commits semanticos (feat:, fix:, docs:, chore:)."
)
h(3, "14.2 Técnica de elicitación")
for item in [
    "Analisis de dominio: observacion de Rappi, Domicilios.com, TaskRabbit, apps municipales.",
    "Modulos diferenciadores: asignacion automatica, aprobacion de prestadores, directorio de emergencias.",
    "Analisis de roles: cuatro perfiles con necesidades distintas.",
    "Adopcion de Clean Architecture y Riverpod para mantenibilidad y escalabilidad.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(4, "14.3 Historias de usuario")
modulos_hu = {
    "MÓDULO 1: Autenticación y Usuarios": [
        (
            "HU-01",
            "Registro",
            "usuario nuevo",
            "registrarme con nombre, apellido, correo y contrasena",
            "crear una cuenta y acceder a funciones segun mi rol",
        ),
        (
            "HU-02",
            "Login",
            "usuario registrado",
            "iniciar sesion con mi correo y contrasena",
            "acceder a mi perfil y funciones segun mi rol",
        ),
        (
            "HU-03",
            "Logout",
            "usuario autenticado",
            "cerrar sesion de forma segura",
            "mis tokens no queden activos en el dispositivo",
        ),
    ],
    "MÓDULO 2: Tienda": [
        (
            "HU-04",
            "Ver tienda",
            "cliente",
            "ver los comercios disponibles y sus productos",
            "elegir que comprar",
        ),
        (
            "HU-05",
            "Carrito",
            "cliente",
            "agregar productos al carrito y modificar cantidades",
            "preparar mi compra antes de confirmarla",
        ),
        (
            "HU-06",
            "Confirmar orden",
            "cliente",
            "confirmar mi carrito y crear una orden",
            "el comercio reciba mi pedido",
        ),
        (
            "HU-07",
            "Historial ordenes",
            "cliente",
            "ver el historial de mis ordenes con sus estados",
            "hacer seguimiento de mis compras",
        ),
    ],
    "MÓDULO 3: Domicilios": [
        (
            "HU-08",
            "Solicitar domicilio",
            "cliente",
            "crear una solicitud de domicilio con direccion de recogida y entrega",
            "que me envien lo que necesito",
        ),
        (
            "HU-09",
            "Estado domicilio",
            "cliente",
            "consultar el estado de mi solicitud en tiempo real",
            "saber cuando llega mi entrega",
        ),
        (
            "HU-10",
            "Gestionar domicilios",
            "domiciliario",
            "ver las solicitudes asignadas y actualizar su estado",
            "gestionar mis entregas",
        ),
        (
            "HU-11",
            "Historial financiero",
            "domiciliario",
            "ver mis ingresos y la comision de Runners",
            "llevar control de mis ganancias",
        ),
    ],
    "MÓDULO 4: Servicios": [
        (
            "HU-12",
            "Ver prestadores",
            "cliente",
            "ver el directorio de prestadores aprobados",
            "contactar al mas adecuado",
        ),
        (
            "HU-13",
            "Solicitar servicio",
            "cliente",
            "enviar una solicitud de servicio con fecha y descripcion",
            "el prestador pueda aceptarla o rechazarla",
        ),
        (
            "HU-14",
            "Perfil prestador",
            "prestador",
            "completar mi perfil con foto, categoria y hoja de vida",
            "el administrador pueda aprobarlo",
        ),
        (
            "HU-15",
            "Gestionar solicitudes",
            "prestador",
            "ver y aprobar/rechazar solicitudes de clientes",
            "gestionar mi agenda de trabajo",
        ),
    ],
    "MÓDULO 5: Contactos": [
        (
            "HU-16",
            "Ver contactos",
            "usuario",
            "ver el directorio de contactos de emergencia y utilidad",
            "encontrar un contacto rapidamente",
        ),
        (
            "HU-17",
            "Filtrar contactos",
            "usuario",
            "filtrar el directorio por tipo de contacto",
            "encontrar mas rapido lo que busco",
        ),
    ],
    "MÓDULO 6: Administración": [
        (
            "HU-18",
            "Dashboard",
            "administrador",
            "ver un panel con metricas de la plataforma",
            "monitorear el estado del sistema",
        ),
        (
            "HU-19",
            "Aprobar prestadores",
            "administrador",
            "revisar perfiles pendientes y aprobar/rechazar",
            "mantener la calidad de los servicios",
        ),
        (
            "HU-20",
            "Gestionar contactos",
            "administrador",
            "crear, editar y desactivar contactos",
            "mantener la informacion publica actualizada",
        ),
        (
            "HU-21",
            "Reportes",
            "administrador",
            "consultar reportes por periodo",
            "tomar decisiones operativas basadas en datos",
        ),
    ],
}
for modulo, hus in modulos_hu.items():
    p(modulo, bold=True)
    for hu_id, nombre, rol, quiero, para in hus:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(
            f"{hu_id} - {nombre}: Como {rol}, quiero {quiero}, para {para}."
        ).font.name = "Calibri"
    doc.add_paragraph()

h(5, "14.4 Diseño relacional")
p("Figura 1. Diseño relacional del sistema Runners. Elaboracion propia (2025).")
p()
entidades = [
    "User: id (PK), email (UK), password, first_name, last_name, role [CLIENTE|PRESTADOR|DOMICILIARIO|ADMIN], is_active, date_joined",
    "StoreCategory: id (PK), name, description",
    "Commerce: id (PK), name, description, city, address, phone, logo, category_id (FK), owner_id (FK), is_active",
    "ProductCategory: id (PK), name",
    "Product: id (PK), name, description, price, commerce_id (FK), category_id (FK), image, is_available",
    "Order: id (PK), client_id (FK), commerce_id (FK), status [PENDIENTE|CONFIRMADO|ENTREGADO|CANCELADO], total_price, notes, created_at",
    "OrderItem: id (PK), order_id (FK), product_id (FK), quantity, unit_price",
    "Deliverer: id (PK), user_id (FK, OneToOne), status [DISPONIBLE|OCUPADO], assigned_number, is_active",
    "DeliveryRequest: id (PK), client_id (FK), deliverer_id (FK, nullable), pickup_address, delivery_address, delivery_fee, status, completed_at, created_at",
    "FinancialRecord: id (PK), deliverer_id (FK), record_type [INGRESO|GASTO], amount, runners_commission, description, created_at",
    "SystemConfig: id (PK), key (UK), value, description",
    "ServiceCategory: id (PK), name, description",
    "ServiceProvider: id (PK), user_id (FK, OneToOne), category_id (FK), description, photo, resume, terms_accepted, status [ACTIVO|INACTIVO], approval_status [PENDIENTE|APROBADO|RECHAZADO], rejection_reason, approved_by_id (FK), approved_at",
    "ServiceRequest: id (PK), client_id (FK), provider_id (FK), category_id (FK), description, scheduled_date, provider_fee, runners_fee, client_total, status [PENDIENTE|APROBADO|RECHAZADO|COMPLETADO], notes, created_at",
    "Contact: id (PK), name, phone, description, contact_type [EMERGENCIA|PROFESIONAL|COMERCIO|OTRO], is_active, created_at, updated_at",
]
for ent in entidades:
    bp = doc.add_paragraph(style="List Bullet")
    run_e = bp.add_run(ent)
    run_e.font.name = "Courier New"
    run_e.font.size = Pt(9)
p()
p(
    "Ver diagrama Mermaid ER completo en Documento_Desarrollo_Runners.md, seccion 14.4.5.",
    italic=True,
)

h(6, "14.4.7 Diagramas de secuencia")
p(
    "Principales diagramas de secuencia del sistema Runners para los flujos criticos de autenticacion, tienda y domicilios."
)

h(6, "14.4.7.1 Autenticacion y enrutamiento por rol")
for item in [
    "El usuario ingresa correo y contrasena en la App Flutter.",
    "La app envia POST /api/v1/auth/login/ a la API Django (Auth).",
    "El backend valida credenciales y retorna tokens JWT (access, refresh) y el rol del usuario.",
    "La app guarda los tokens en FlutterSecureStorage.",
    "La capa de enrutamiento (GoRouter) recibe el rol autenticado.",
    "GoRouter redirige a la Shell correspondiente (cliente, domiciliario, prestador, admin).",
    "El usuario visualiza la pantalla inicial propia de su rol.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.2 Compra en tienda y creacion de orden")
for item in [
    "El cliente navega por la tienda y agrega productos al carrito.",
    "La app actualiza el carrito local usando Hive (persistencia offline).",
    "Al confirmar el carrito, la app construye el payload de la orden.",
    "Se envia POST /api/v1/store/orders/ con los productos y cantidades.",
    "El backend crea la Order y los OrderItems en la base de datos.",
    "La API responde 201 Created con el detalle de la orden y su estado inicial PENDIENTE.",
    "La app muestra al cliente el resumen de compra y el estado de la orden.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.3 Solicitud de domicilio con asignacion automatica")
for item in [
    "El cliente completa el formulario de domicilio (origen, destino, descripcion).",
    "La app envia POST /api/v1/deliveries/requests/ a la API de domicilios.",
    "El backend busca un domiciliario con estado DISPONIBLE y menor assigned_number.",
    "Si encuentra uno, crea un DeliveryRequest asociado al domiciliario; de lo contrario, deja deliverer = null.",
    "La API responde 201 Created con la solicitud y el estado inicial (ACEPTADO o PENDIENTE).",
    "La app muestra al cliente el detalle y el estado del domicilio.",
    "El domiciliario ve la solicitud asignada y actualiza su estado (EN_CAMINO, ENTREGADO), generando registros financieros.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.4 Registro de usuario con emision de tokens")
for item in [
    "El usuario diligencia el formulario de registro en la App Flutter.",
    "La app envia POST /api/v1/auth/register/ a la API de usuarios.",
    "El backend valida los datos y crea un nuevo User en la base de datos.",
    "Se generan automaticamente los tokens JWT (access y refresh) para el nuevo usuario.",
    "La API responde 201 Created con los datos del usuario y los tokens.",
    "La app guarda los tokens en FlutterSecureStorage para sesiones posteriores.",
    "La navegacion redirige al flujo inicial segun el rol (CLIENTE por defecto o PRESTADOR en proceso de aprobacion).",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.5 Registro y aprobacion de prestador de servicios")
for item in [
    "Un usuario autenticado completa su perfil como prestador (foto, descripcion, categoria, HV).",
    "La app envia POST /api/v1/services/providers/register/ a la API de servicios.",
    "El backend crea el objeto ServiceProvider con approval_status = PENDIENTE.",
    "Se actualiza el rol del usuario a PRESTADOR manteniendo sus credenciales.",
    "El administrador consulta la lista de prestadores pendientes desde el dashboard.",
    "La app envia la decision (approve o reject) a un endpoint de aprobacion de prestadores.",
    "El backend actualiza approval_status, status y los campos de auditoria (approved_by, approved_at, rejection_reason).",
    "El prestador es notificado en la app sobre el resultado de la revision.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.6 Solicitud de servicio entre cliente y prestador")
for item in [
    "El cliente navega por el directorio de prestadores aprobados desde la app.",
    "La app consulta la API de servicios filtrando prestadores con approval_status = APROBADO.",
    "El cliente selecciona un prestador, define fecha y descripcion del servicio requerido.",
    "Se envia POST /api/v1/services/requests/ para crear la solicitud de servicio.",
    "El backend registra un ServiceRequest asociado al cliente y al prestador con un estado inicial (por ejemplo, REGISTRADA).",
    "El cliente visualiza la confirmacion y el estado de su solicitud en la app.",
    "El prestador consulta sus solicitudes desde la app (modo PRESTADOR) usando la lista de ServiceRequest asignados.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.4.7.7 Consulta de contactos y dashboard administrador")
for item in [
    "Cualquier usuario abre el modulo de contactos en la app.",
    "La app consulta GET /api/v1/contacts/ aplicando filtros por tipo de contacto cuando corresponde.",
    "La API retorna solo los contactos activos, que se muestran en la interfaz con opciones rapidas de llamada.",
    "El administrador ingresa al dashboard administrativo desde su Shell.",
    "La app llama al endpoint de resumen (dashboard) para obtener metricas de usuarios, ordenes, domicilios y servicios.",
    "La informacion recibida se presenta en tarjetas y graficos que resumen el estado operativo de la plataforma.",
]:
    bp = doc.add_paragraph(style="List Number")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.5 Prototipado")
p(
    "Prototipo funcional en Flutter que incluye pantallas de login/registro, home por modulos segun rol, tienda con carrito, flujo de domicilios, directorio de servicios y contactos, y panel de administracion."
)

h(6, "14.6 Plan de pruebas y aseguramiento de la calidad")
p(
    "Plan de pruebas del proyecto Runners para validar requisitos funcionales y no funcionales en todos los modulos."
)

h(6, "14.6.1 Objetivos de las pruebas")
for item in [
    "Verificar el cumplimiento de los requisitos funcionales y no funcionales descritos para cada modulo.",
    "Detectar defectos de forma temprana durante el ciclo de desarrollo para reducir el costo de correccion.",
    "Asegurar la estabilidad de los flujos criticos: autenticacion, tienda, domicilios, servicios y administracion.",
    "Validar que el sistema sea utilizable y comprensible para CLIENTE, PRESTADOR, DOMICILIARIO y ADMIN.",
    "Proveer evidencia documentada del funcionamiento correcto del sistema para la evaluacion academica.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.6.2 Alcance y niveles de prueba")
p(
    "Se consideran pruebas unitarias (parciales), de integracion, de sistema end-to-end y de aceptacion academica."
)
for item in [
    "Pruebas unitarias parciales en backend (logica de negocio en modelos y funciones auxiliares) y en frontend (widgets y validacion de formularios).",
    "Pruebas de integracion entre Flutter y Django usando Dio y la API REST (serializacion/deserializacion correcta).",
    "Pruebas de sistema (end-to-end) que recorren flujos completos por rol, incluyendo escenarios de exito y error.",
    "Pruebas de aceptacion academica mediante demostraciones guiadas de los casos de uso principales.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.6.3 Estrategia de pruebas por modulo")
tbl_tp = doc.add_table(rows=1, cols=4)
tbl_tp.style = "Table Grid"
hdr_tp = tbl_tp.rows[0].cells
hdr_tp[0].text = "Modulo"
hdr_tp[1].text = "Tipo de pruebas"
hdr_tp[2].text = "Herramientas"
hdr_tp[3].text = "Descripcion"
estrategia = [
    ("Autenticacion y usuarios", "Integracion, sistema", "Postman, App Flutter, Django Admin", "Registro, login, logout, perfil y restricciones por rol."),
    ("Tienda", "Integracion, sistema", "Postman, App Flutter", "Listado de comercios/productos, carrito y ordenes."),
    ("Domicilios", "Integracion, sistema", "Postman, App Flutter", "Solicitudes, autoasignacion, estados y registros financieros."),
    ("Servicios", "Integracion, sistema", "Postman, App Flutter, Django Admin", "Registro de prestadores, aprobacion/rechazo y ciclo de solicitudes."),
    ("Contactos", "Sistema", "Postman, App Flutter", "Directorio publico filtrado y gestion por administrador."),
    ("Administracion y reportes", "Sistema", "Postman, App Flutter", "Dashboard, reportes de ventas, domiciliarios y servicios."),
]
for modulo, tipo_p, herramientas, desc in estrategia:
    row = tbl_tp.add_row().cells
    row[0].text = modulo
    row[1].text = tipo_p
    row[2].text = herramientas
    row[3].text = desc
doc.add_paragraph()

h(6, "14.6.4 Casos de prueba representativos")
casos_tp = [
    "PT-01 – Registro exitoso de usuario (HU-01, UC-01): crear usuario, generar tokens y redirigir segun rol.",
    "PT-02 – Inicio de sesion con credenciales invalidas (HU-02, UC-02): respuesta 401 y mensaje de error amigable.",
    "PT-03 – Flujo completo de compra en tienda (HU-04 a HU-07): listar, agregar al carrito, confirmar y verificar historial.",
    "PT-04 – Solicitud de domicilio sin domiciliarios disponibles (HU-08): mensaje de indisponibilidad y sin creacion de registro.",
    "PT-05 – Aprobacion de prestador y visibilidad en directorio (HU-14, HU-19): cambio de estado y aparicion en listado publico.",
    "PT-06 – Directorio de contactos filtrado (HU-16, HU-17): filtros por tipo EMERGENCIA, COMERCIO, PROFESIONAL.",
    "PT-07 – Dashboard administrativo coherente (HU-18, HU-21): metricas y totales consistentes con los datos en BD.",
]
for item in casos_tp:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.6.5 Entorno de pruebas")
for item in [
    "Backend: Django 5.2.6 + DRF 3.16.1 sobre Python 3.12 con base de datos SQLite (db.sqlite3) y datos semilla.",
    "Frontend: App Flutter en emulador Android o dispositivo fisico, con archivo .env apuntando al backend local.",
    "Herramientas de apoyo: Postman (coleccion runners_api_postman.json) y Django Admin para inspeccion rapida durante pruebas.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(6, "14.6.6 Criterios de aceptacion y salida")
for item in [
    "Todos los casos de prueba criticos (PT-01 a PT-07) ejecutados sin errores bloqueantes.",
    "Sin errores de severidad alta abiertos en autenticacion, tienda, domicilios o servicios al finalizar el ciclo de pruebas.",
    "Ejecucion completa y coherente de los casos de uso UC-01 a UC-13 de principio a fin.",
    "Ejecucion de pruebas de regresion basicas tras cambios significativos en modelos/serializers y providers principales.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

p("La evidencia de ejecucion incluye capturas de pantalla, videos de flujos principales y registros de llamadas en Postman.")

# ===================== 15. CASOS DE USO =====================
h(1, "15. Casos de uso (detallados)")
p("Actores: Cliente, Prestador, Domiciliario, Administrador, Sistema (API Django).")
casos = [
    (
        "UC-01",
        "Registrarse",
        "Usuario no registrado",
        "POST /api/auth/register/ - valida correo unico - crea usuario - retorna tokens JWT",
        "Email ya registrado, contrasena muy corta, rol invalido",
    ),
    (
        "UC-02",
        "Iniciar sesion",
        "Cuenta existente",
        "POST /api/auth/login/ - valida credenciales - retorna tokens - GoRouter redirige segun rol",
        "Credenciales invalidas, cuenta inactiva",
    ),
    (
        "UC-03",
        "Consultar tienda",
        "-",
        "GET /api/store/commerces/ - GET /api/store/products/?commerce_id=X",
        "-",
    ),
    (
        "UC-04",
        "Gestionar carrito",
        "-",
        "CartNotifier actualiza Hive - POST /api/store/orders/ al confirmar",
        "Carrito vacio",
    ),
    (
        "UC-05",
        "Solicitar domicilio",
        "-",
        "POST /api/deliveries/requests/ - asignacion automatica de domiciliario disponible",
        "Sin domiciliarios disponibles",
    ),
    (
        "UC-06",
        "Actualizar estado domicilio",
        "-",
        "PATCH /api/deliveries/requests/{id}/ - si ENTREGADO: completed_at + FinancialRecord",
        "-",
    ),
    (
        "UC-07",
        "Completar perfil prestador",
        "-",
        "POST /api/services/providers/ con foto, descripcion, categoria, HV - approval_status = PENDIENTE",
        "-",
    ),
    (
        "UC-08",
        "Aprobar/rechazar prestador",
        "Admin",
        "PATCH /api/services/providers/{id}/ con approval_status = APROBADO o RECHAZADO",
        "-",
    ),
    (
        "UC-09",
        "Solicitar servicio",
        "Cliente",
        "POST /api/services/requests/ - sistema calcula client_total - prestador recibe solicitud",
        "Prestador no aprobado, solicitud duplicada",
    ),
    (
        "UC-10",
        "Gestionar solicitudes servicio",
        "Prestador",
        "GET /api/services/requests/ - PATCH /{id}/ con status APROBADO o RECHAZADO",
        "-",
    ),
    (
        "UC-11",
        "Consultar contactos",
        "-",
        "GET /api/contacts/?contact_type=X - Flutter muestra directorio",
        "-",
    ),
    (
        "UC-12",
        "Administrar plataforma",
        "Admin",
        "Endpoints de usuarios, comercios, configuracion y contactos",
        "-",
    ),
    (
        "UC-13",
        "Ver reportes",
        "Admin",
        "GET /api/reports/deliveries/ + /services/ + /store/ con filtros de periodo",
        "-",
    ),
]
for uc_id, nombre, pre, flujo, exc in casos:
    h(2, f"{uc_id}: {nombre}")
    if pre and pre != "-":
        p(f"Precondicion: {pre}")
    p(f"Flujo: {flujo}")
    if exc and exc != "-":
        p(f"Excepciones: {exc}")

h(3, "15.1 Diagrama de casos de uso")
p(
    "Ver diagrama Mermaid en Documento_Desarrollo_Runners.md, seccion 15.1.",
    italic=True,
)

# ===================== 16. ESTRUCTURA =====================
h(1, "16. Estructura de carpetas")
h(3, "16.1 Backend (Django)")
codigo_backend = """backend/
+-- manage.py
+-- requirements.txt
+-- .env / .env.example
+-- runners_project/
|   +-- settings.py
|   +-- urls.py
|   +-- wsgi.py
+-- apps/
|   +-- users/
|   +-- store/
|   +-- deliveries/
|   +-- services/
|   +-- contacts/
|   +-- reports/
+-- media/"""
code(codigo_backend)

h(3, "16.2 Frontend (Flutter + Clean Architecture)")
codigo_flutter = """frontend/
+-- pubspec.yaml
+-- lib/
|   +-- main.dart
|   +-- app.dart
|   +-- core/
|   |   +-- network/    # Dio + interceptores JWT
|   |   +-- router/     # GoRouter con guardias de rol
|   |   +-- storage/    # Hive + SecureStorage
|   |   +-- theme/      # AppTheme
|   +-- features/
|   |   +-- auth/
|   |   +-- store/
|   |   +-- deliveries/
|   |   +-- services/
|   |   +-- contacts/
|   |   +-- admin/
|   +-- shared/
|       +-- widgets/    # AppButton, AppTextField, AppCard
+-- test/"""
code(codigo_flutter)

# ===================== 17-22. RIESGOS =====================
h(1, "17. Plan de Gestión de Riesgos")
h(2, "17.1 Introduccion")
p(
    "Esta seccion establece estrategias, procedimientos y responsabilidades para identificar, analizar, mitigar y controlar los riesgos del proyecto Runners a lo largo de su ciclo de vida."
)
h(3, "17.1.1 Proposito")
p(
    "Definir la metodologia y acciones para gestionar los riesgos del proyecto Runners, garantizando la entrega a tiempo de una plataforma funcional, segura y documentada."
)
h(4, "17.1.2 Alcance")
p(
    "Abarca todas las fases: analisis, diseno, desarrollo, pruebas y despliegue. Aplica a todo el equipo, herramientas y procesos de gestion interna."
)
h(5, "17.1.3 Definiciones y Acronimos")
for item in [
    "DRF: Django REST Framework",
    "JWT: JSON Web Token",
    "HU: Historia de Usuario",
    "RN: Regla de Negocio",
    "UC: Caso de Uso",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"
h(6, "17.1.5 Descripcion General")
for item in [
    "Resumen de riesgos: lista priorizada.",
    "Tareas de gestion: identificacion, analisis y mitigacion.",
    "Responsabilidades: roles del equipo.",
    "Presupuesto: tiempo para contingencias.",
    "Herramientas: Git, GitHub Issues, documentacion colaborativa.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(1, "18. Tareas de Gestión de Riesgos")
for item in [
    "Identificacion: analisis con el equipo revisando cada modulo.",
    "Analisis: Matriz de probabilidad/impacto (escala 1-5).",
    "Priorizacion: Riesgos con puntuacion >= 8 se gestionan primero.",
    "Seguimiento: revision semanal durante el sprint activo.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(1, "19. Organización y Responsabilidades")
for item in [
    "Lider del Proyecto: coordina actividades, gestiona comunicacion y toma decisiones de alcance.",
    "Gestor de Riesgos: identifica, documenta y hace seguimiento a riesgos.",
    "Equipo Tecnico (todos): deteccion temprana de riesgos tecnicos.",
    "Comite de Revision: revisiones periodicas para evaluar riesgos e implementar medidas correctivas.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(1, "20. Presupuesto")
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
tbl3.rows[0].cells[0].text = "Recurso"
tbl3.rows[0].cells[1].text = "Costo"
for recurso, costo in [
    ("Servidor de desarrollo (local)", "$0"),
    ("Repositorio GitHub (publico)", "$0"),
    ("Herramientas (VS Code, Postman)", "$0"),
    ("Flutter SDK, Django, librerias open source", "$0"),
    ("Servidor de produccion (estimado futuro)", "Variable"),
    ("Contingencia de tiempo (15% del sprint)", "4-6 horas por sprint"),
]:
    row = tbl3.add_row().cells
    row[0].text = recurso
    row[1].text = costo
doc.add_paragraph()

h(1, "21. Herramientas y Técnicas")
for item in [
    "Control de versiones: Git + GitHub (ramas feature, commits semanticos, Pull Requests).",
    "Gestion de tareas: GitHub Issues / tablero Kanban.",
    "Documentacion: Markdown en el repositorio (README, QUICKSTART, CONTRIBUTING).",
    "Pruebas de API: Postman con coleccion runners_api_postman.json (30+ endpoints).",
    "Pruebas de Frontend: flujos manuales completos por modulo.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

h(1, "22. Elementos de Riesgo Por Gestionar")
riesgos = [
    (
        "Riesgo 1 – Perdida/ausencia de integrante",
        "Documentacion en GitHub, arquitectura modular, roles redundantes.",
    ),
    (
        "Riesgo 2 – Perdida o dano de archivos",
        "Todo versionado en GitHub; push frecuente; clone en cualquier equipo.",
    ),
    (
        "Riesgo 3 – Falta de comunicacion",
        "Reglas de branches, commits semanticos, PR con descripcion, reuniones periodicas.",
    ),
    (
        "Riesgo 4 – Incumplimiento de plazos",
        "Buffer 15% por sprint, reducir alcance si es necesario.",
    ),
    (
        "Riesgo 5 – Problemas con dependencias",
        "Versiones fijas en pubspec.yaml y requirements.txt; entorno virtual Python aislado.",
    ),
    (
        "Riesgo 6 – Falta de experiencia con el stack",
        "Documentacion de arquitectura, revision entre pares, guias de implementacion en el repo.",
    ),
    (
        "Riesgo 7 – Cambios en los requisitos",
        "Arquitectura modular; migraciones Django para cambios en el modelo.",
    ),
    (
        "Riesgo 8 – Fallas en despliegue a produccion",
        "Variables de entorno en .env; documentacion en QUICKSTART.md.",
    ),
]
for titulo, control in riesgos:
    p(titulo, bold=True)
    p(control, indent=True)
    doc.add_paragraph()

p("Niveles de riesgo: Bajo (1-4) | Medio (5-9) | Alto (10-14) | Muy alto (15-25)")

# ===================== 23. DIAGRAMA DE PAQUETES =====================
h(1, "23. Diagrama de paquetes")
p(
    "Ver diagrama Mermaid en Documento_Desarrollo_Runners.md, seccion 23.",
    italic=True,
)

# ===================== 24. CRONOGRAMA =====================
h(1, "24. Cronograma")
h(2, "Planificacion del tiempo")
p("Duracion del Sprint: 2 semanas (10 dias habiles)")
p(
    "Capacidad: 3 desarrolladores x 40 hr/semana x 2 semanas = 232 hr efectivas = 29 puntos de historia"
)
p()
p("Sprint 1 – Autenticacion + Tienda (21 puntos)", bold=True)
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = "Table Grid"
tbl4.rows[0].cells[0].text = "Historia"
tbl4.rows[0].cells[1].text = "Tareas"
tbl4.rows[0].cells[2].text = "Pts"
for historia, tareas, pts in [
    ("HU-01 Registro", "Pantalla registro, validacion correo unico, serializer", "2"),
    ("HU-02 Login", "Pantalla login, endpoint JWT, interceptor Dio, SecureStorage", "3"),
    ("HU-03 Logout", "Limpiar tokens, redirigir a Login", "1"),
    ("HU-04 Ver tienda", "Endpoint comercios/productos, pantalla listado", "4"),
    ("HU-05 Carrito", "CartNotifier Riverpod, persistencia Hive, pantalla carrito", "4"),
    ("HU-06 Confirmar orden", "Endpoint crear orden con items, pantalla confirmacion", "4"),
    ("HU-07 Historial", "Endpoint ordenes del cliente, pantalla historial", "3"),
]:
    row = tbl4.add_row().cells
    row[0].text = historia
    row[1].text = tareas
    row[2].text = pts
doc.add_paragraph()

p("Sprint 2 – Domicilios + Servicios (23 puntos)", bold=True)
tbl5 = doc.add_table(rows=1, cols=3)
tbl5.style = "Table Grid"
tbl5.rows[0].cells[0].text = "Historia"
tbl5.rows[0].cells[1].text = "Tareas"
tbl5.rows[0].cells[2].text = "Pts"
for historia, tareas, pts in [
    ("HU-08 Solicitar domicilio", "Endpoint + asignacion automatica domiciliario", "4"),
    ("HU-09 Estado domicilio", "Pantalla seguimiento, endpoint detalle", "2"),
    ("HU-10 Gestionar domicilios", "Pantalla domiciliario, PATCH estado, FinancialRecord", "4"),
    ("HU-11 Historial financiero", "Endpoint registros financieros, pantalla", "2"),
    ("HU-12 Ver prestadores", "Endpoint prestadores aprobados, pantalla directorio", "3"),
    ("HU-13 Solicitar servicio", "Endpoint solicitud, calculo tarifas", "3"),
    ("HU-14 Perfil prestador", "Endpoint perfil, subida foto y HV", "3"),
    ("HU-15 Gestionar solicitudes", "Aprobar/rechazar solicitudes, pantalla prestador", "2"),
]:
    row = tbl5.add_row().cells
    row[0].text = historia
    row[1].text = tareas
    row[2].text = pts
doc.add_paragraph()

p("Sprint 3 – Contactos + Admin + Refinamiento (21 puntos)", bold=True)
tbl6 = doc.add_table(rows=1, cols=3)
tbl6.style = "Table Grid"
tbl6.rows[0].cells[0].text = "Historia"
tbl6.rows[0].cells[1].text = "Tareas"
tbl6.rows[0].cells[2].text = "Pts"
for historia, tareas, pts in [
    ("HU-16 Ver contactos", "Endpoint directorio, pantalla listado filtrable", "2"),
    ("HU-17 Filtrar contactos", "Filtros por tipo en endpoint y UI", "1"),
    ("HU-18 Dashboard", "Endpoint metricas, pantalla con cards", "4"),
    ("HU-19 Aprobar prestadores", "Endpoint PATCH approval, pantalla gestion", "3"),
    ("HU-20 Gestionar contactos", "CRUD contactos para admin, pantalla gestion", "3"),
    ("HU-21 Reportes", "Endpoints reportes, pantallas de consulta", "4"),
    ("Buffer / pruebas", "Pruebas Postman, flujos completos, ajustes UX", "4"),
]:
    row = tbl6.add_row().cells
    row[0].text = historia
    row[1].text = tareas
    row[2].text = pts
doc.add_paragraph()

# ===================== 25. RESULTADOS ESPERADOS =====================
h(1, "25. Resultados Esperados de la Aplicacion")
for item in [
    "Digitalizacion de la economia local: conectar clientes con comercios y prestadores en un canal unico.",
    "Eficiencia en la asignacion de domicilios: asignacion automatica por disponibilidad.",
    "Validacion de prestadores: calidad garantizada mediante aprobacion administrativa previa.",
    "Centralizacion de informacion: visibilidad completa para el administrador en tiempo real.",
    "Reduccion de errores: flujos automatizados minimizan la intervencion manual.",
    "Arquitectura escalable: Clean Architecture + DRF permite agregar modulos sin reestructurar.",
    "Experiencia de usuario fluida: Riverpod, Hive y manejo de errores garantizan confiabilidad.",
    "Documentacion completa: README, QUICKSTART, guias, Postman y este documento.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 26. RECOMENDACIONES =====================
h(1, "26. Recomendaciones y Sugerencias a Futuro")
for titulo, desc in [
    (
        "1. Pagos en linea reales",
        "Integrar PayPal, Stripe o MercadoPago. Requiere adecuarse a normativa de pagos electronicos.",
    ),
    (
        "2. Notificaciones push",
        "Firebase Cloud Messaging (FCM) para cambios de estado en domicilios y nuevas solicitudes.",
    ),
    (
        "3. Chat integrado",
        "Django Channels + WebSockets en Flutter para comunicacion cliente-prestador/domiciliario.",
    ),
    (
        "4. Sistema de calificaciones",
        "Rating para domiciliarios y prestadores despues de completar un servicio.",
    ),
    (
        "5. Geolocalizacion",
        "GPS para seguimiento de domicilios en mapa (Google Maps SDK para Flutter).",
    ),
    (
        "6. Analytics avanzado",
        "Graficos de ventas, servicios mas solicitados, domiciliarios mas activos.",
    ),
    (
        "7. Aplicacion web",
        "Con la API ya desarrollada, construir PWA o React para administradores y comercios.",
    ),
    (
        "8. Docker y produccion",
        "Docker Compose con Django, PostgreSQL, Nginx, SSL y gunicorn para despliegue robusto.",
    ),
    (
        "9. IA y recomendaciones",
        "Recomendar prestadores/productos basados en historial usando filtrado colaborativo.",
    ),
    (
        "10. Auditoria y trazabilidad",
        "Log de acciones administrativas para mejorar seguridad y control.",
    ),
]:
    h(3, titulo)
    p(desc)

# ===================== 27. CONCLUSIONES =====================
h(1, "27. Conclusiones")
for item in [
    "La arquitectura propuesta (Flutter + Clean Architecture + Riverpod; Django + DRF + JWT) ofrece una base solida, productiva y mantenible.",
    "El modelo relacional disenado soporta correctamente los cuatro modulos con sus relaciones entre entidades.",
    "La asignacion automatica de domiciliarios simplifica el flujo operativo sin intervencion manual.",
    "El sistema de aprobacion de prestadores garantiza calidad en el directorio de servicios.",
    "Clean Architecture en Flutter asegura que cambios en la capa de datos no afecten la presentacion.",
    "Para produccion: PostgreSQL, almacenamiento S3, HTTPS, Docker y CI/CD en GitHub Actions.",
    "Se recomienda implementar notificaciones push (FCM) y geolocalizacion como proximas funcionalidades.",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(item).font.name = "Calibri"

# ===================== 28. BIBLIOGRAFÍA =====================
h(1, "28. Bibliografia")
for ref in [
    "Sommerville, I. (2016). Ingenieria del Software. Pearson.",
    "Pressman, R. (2015). Ingenieria del Software: Un enfoque practico. McGraw-Hill.",
    "Django Software Foundation. Django Documentation. https://docs.djangoproject.com/",
    "Django REST Framework. DRF Documentation. https://www.django-rest-framework.org/",
    "JWT.io. JSON Web Tokens Introduction. https://jwt.io/introduction",
    "Flutter Team. Flutter Documentation. https://docs.flutter.dev/",
    "Riverpod. Riverpod Documentation. https://riverpod.dev/",
    "GoRouter. go_router package. https://pub.dev/packages/go_router",
    "Hive. Hive Documentation. https://docs.hivedb.dev/",
    "Simple JWT. djangorestframework-simplejwt Docs. https://django-rest-framework-simplejwt.readthedocs.io/",
    "Mermaid. Diagramming Tool. https://mermaid.js.org/",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run(ref).font.name = "Calibri"

doc.save("Documento_Desarrollo_Runners.docx")
print("DOCX generado exitosamente: Documento_Desarrollo_Runners.docx")
