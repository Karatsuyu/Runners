"""
Convierte Documento_de_Desarrollo_Runners.md a .docx
respetando todo el contenido tal cual, con formato básico.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = Path(__file__).parent / "Documento_de_Desarrollo_Runners.md"
DOCX_PATH = Path(__file__).parent / "Documento_de_Desarrollo_Runners.docx"


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_run(paragraph, text, bold=False, italic=False, font_size=None, color=None, font_name=None):
    """Add a run with optional formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    return run


def parse_inline(paragraph, text, base_bold=False, base_italic=False, font_size=None, font_name=None):
    """Parse inline markdown formatting (**bold**, *italic*, `code`) and add runs."""
    # Pattern to match **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)'
    last_end = 0
    for m in re.finditer(pattern, text):
        # Add text before match
        if m.start() > last_end:
            add_formatted_run(paragraph, text[last_end:m.start()], bold=base_bold, italic=base_italic, font_size=font_size, font_name=font_name)
        if m.group(2):  # **bold**
            add_formatted_run(paragraph, m.group(2), bold=True, italic=base_italic, font_size=font_size, font_name=font_name)
        elif m.group(3):  # *italic*
            add_formatted_run(paragraph, m.group(3), bold=base_bold, italic=True, font_size=font_size, font_name=font_name)
        elif m.group(4):  # `code`
            run = add_formatted_run(paragraph, m.group(4), bold=base_bold, italic=base_italic, font_size=font_size if font_size else None, font_name='Consolas')
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        last_end = m.end()
    # Remaining text
    if last_end < len(text):
        add_formatted_run(paragraph, text[last_end:], bold=base_bold, italic=base_italic, font_size=font_size, font_name=font_name)


def convert():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # --- Code block ---
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block - flush
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.left_indent = Cm(1)
                for ci, cl in enumerate(code_lines):
                    run = p.add_run(cl)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    if ci < len(code_lines) - 1:
                        p.add_run("\n")
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Table detection ---
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            # Check if next line is separator
            if not in_table:
                in_table = True
                table_rows = []
            # Parse cells
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows like |---|---|
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is still table
            next_i = i + 1
            if next_i < len(lines):
                nl = lines[next_i].strip()
                if not (nl.startswith("|") and "|" in nl[1:]):
                    # End of table, flush
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
            else:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        if in_table:
            # Non-table line reached; flush
            _flush_table(doc, table_rows)
            table_rows = []
            in_table = False
            # Don't increment i, re-process this line
            continue

        stripped = line.strip()

        # --- Horizontal rule ---
        if stripped == "---":
            # Add a subtle horizontal line paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # We just add a thin line via bottom border on the paragraph
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            heading_level = min(level, 4)  # docx supports Heading 1-9
            h = doc.add_heading(level=heading_level)
            parse_inline(h, text)
            i += 1
            continue

        # --- Empty line ---
        if stripped == "":
            i += 1
            continue

        # --- Blockquote ---
        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            parse_inline(p, text, base_italic=True)
            # Add left border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="4472C4"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # --- Unordered list ---
        list_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if list_match:
            indent_spaces = len(list_match.group(1))
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Cm(2.0)
            else:
                p.paragraph_format.left_indent = Cm(1.27)
            parse_inline(p, text)
            i += 1
            continue

        # --- Ordered list ---
        ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, text)
            i += 1
            continue

        # --- Normal paragraph ---
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    # Flush remaining table if any
    if in_table and table_rows:
        _flush_table(doc, table_rows)

    doc.save(str(DOCX_PATH))
    print(f"✅ Documento generado: {DOCX_PATH}")


def _flush_table(doc, rows):
    """Create a Word table from parsed rows."""
    if not rows:
        return

    # Determine max columns
    max_cols = max(len(r) for r in rows)
    # Normalize rows
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ri == 0:
                # Header row
                set_cell_shading(cell, "1F3864")
                parse_inline(p, cell_text, base_bold=True, font_size=10, font_name='Calibri')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                parse_inline(p, cell_text, font_size=10, font_name='Calibri')

    # Auto-fit hint
    table.autofit = True
    doc.add_paragraph()  # spacing after table


if __name__ == "__main__":
    convert()
